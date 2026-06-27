"""
fill_contract.py — Remplit le template DOCX du contrat F5L.

Usage:
    from fill_contract import fill_contract
    docx_bytes = fill_contract("/path/to/template.docx", variables_dict)

Toutes les substitutions sont réalisées sans toucher au style existant des cellules.
"""

from __future__ import annotations

import io
from copy import deepcopy
from datetime import date, timedelta
from typing import Any

from docx import Document
from docx.oxml.ns import qn


# ─── Helpers ──────────────────────────────────────────────────────────────────

def _set_cell_text(cell, text: str) -> None:
    """Remplace le texte d'une cellule en préservant le style (font, couleur, taille)."""
    for para in cell.paragraphs:
        for run in para.runs:
            run.text = ""
    if not cell.paragraphs:
        return
    para = cell.paragraphs[0]
    if para.runs:
        para.runs[0].text = str(text)
    else:
        para.add_run(str(text))


def _replace_in_para(para, old: str, new: str) -> None:
    """Remplace une sous-chaîne dans un paragraphe en préservant les runs."""
    full_text = "".join(run.text for run in para.runs)
    if old not in full_text:
        return
    new_text = full_text.replace(old, new)
    if para.runs:
        para.runs[0].text = new_text
        for run in para.runs[1:]:
            run.text = ""
    else:
        para.add_run(new_text)


def _replace_global(doc: Document, old: str, new: str) -> None:
    """Remplace dans tous les paragraphes du doc (corps + tables)."""
    for para in doc.paragraphs:
        _replace_in_para(para, old, new)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_para(para, old, new)


def _check_service(cell, checked: bool) -> None:
    """Coche ou décoche la case d'un service."""
    _set_cell_text(cell, "☑" if checked else "□")


def _add_months(dt: date, months: int) -> date:
    month = dt.month - 1 + months
    year = dt.year + month // 12
    month = month % 12 + 1
    import calendar
    day = min(dt.day, calendar.monthrange(year, month)[1])
    return date(year, month, day)


# ─── Prix parsing ──────────────────────────────────────────────────────────────

def _parse_price(price_str: str) -> int:
    """Extrait la valeur numérique d'un prix (ex: '299 €' → 299)."""
    digits = "".join(c for c in price_str if c.isdigit())
    return int(digits) if digits else 0


# ─── Remplissage principal ────────────────────────────────────────────────────

def fill_contract(template_path: str, v: dict[str, Any]) -> bytes:
    """
    Remplit le template DOCX avec les variables v.
    Retourne les bytes du document rempli.
    """
    doc = Document(template_path)
    tables = doc.tables

    services = v.get("services_souscrits", [])

    # ── 1. TABLE_CLIENT (index 2) — cellules droites ────────────────────────
    if len(tables) > 2:
        t = tables[2]
        fields = [
            "client_nom",
            "client_forme_juridique",
            "client_siret",
            "client_adresse",
            "client_representant",
            "client_email",
            "client_telephone",
        ]
        for i, field in enumerate(fields):
            if i < len(t.rows):
                _set_cell_text(t.rows[i].cells[-1], v.get(field, ""))

    # ── 2. "Fait à..." paragraph ─────────────────────────────────────────────
    lieu = v.get("lieu_signature", "")
    date_sig = v.get("date_signature", "")
    _replace_global(doc, "Fait à _____________________ , le _______ / _______ / 2025",
                    f"Fait à {lieu} , le {date_sig}")

    # ── 3. TABLE_SERVICES (index 5) — services + prix ────────────────────────
    prix_map = {
        "site":     v.get("prix_site", ""),
        "seo":      v.get("prix_seo", ""),
        "ads":      v.get("prix_ads", ""),
        "fidelite": v.get("prix_fidelite", ""),
        "ia":       v.get("prix_ia", ""),
        "social":   v.get("prix_social", ""),
    }
    service_keys = ["site", "seo", "ads", "fidelite", "ia", "social"]

    if len(tables) > 5:
        t = tables[5]
        for i, key in enumerate(service_keys, start=1):
            if i < len(t.rows):
                row = t.rows[i]
                checked = key in services
                _check_service(row.cells[0], checked)
                _set_cell_text(row.cells[3], prix_map[key] if checked else "—")

        # Total
        if len(t.rows) > 7:
            total = v.get("total_mensuel", "")
            _set_cell_text(t.rows[7].cells[-1], total)

    # ── 4. TABLE_DUREE (index 6) — cocher la durée choisie ───────────────────
    duree_labels = {"1": "1 mois", "3": "3 mois", "6": "6 mois", "12": "12 mois"}
    duree_val = str(v.get("duree_mois", "1"))

    if len(tables) > 6:
        t = tables[6]
        for i, key in enumerate(["1", "3", "6", "12"]):
            if i < len(t.rows):
                _check_service(t.rows[i].cells[0], key == duree_val)

    # ── 5. Paragraphes durée + dates ─────────────────────────────────────────
    date_debut = v.get("date_debut", "")
    date_fin   = v.get("date_fin", "")

    _replace_global(doc,
        "Durée d'engagement sélectionnée : ☐ 1 mois   ☐ 3 mois   ☐ 6 mois   ☐ 12 mois",
        "Durée d'engagement sélectionnée : "
        + " ".join(
            ("☑" if k == duree_val else "☐") + f" {duree_labels[k]}"
            for k in ["1", "3", "6", "12"]
        )
    )
    _replace_global(doc,
        "Date de début de la prestation : _______ / _______ / _______",
        f"Date de début de la prestation : {date_debut}"
    )
    _replace_global(doc,
        "Date de fin de l'engagement initial (si applicable) : _______ / _______ / _______",
        f"Date de fin de l'engagement initial (si applicable) : {date_fin}"
    )

    # ── 6. Mode facturation ───────────────────────────────────────────────────
    mode = v.get("mode_facturation", "debut")
    debut_checked = "☑" if mode == "debut" else "☐"
    fin_checked   = "☑" if mode == "fin"   else "☐"
    _replace_global(doc,
        "☐ Facturation en début de mois (mois à venir)   ☐ Facturation en fin de mois (mois écoulé)",
        f"{debut_checked} Facturation en début de mois (mois à venir)   {fin_checked} Facturation en fin de mois (mois écoulé)"
    )

    # ── 7. Email RGPD ────────────────────────────────────────────────────────
    _replace_global(doc, "contact@f5l-agency.fr", "Ladouceurmc.contact@gmail.com")

    # ── 8. TABLE_SIGNATURES (index 13) — bloc client ─────────────────────────
    if len(tables) > 13:
        t = tables[13]
        if t.rows and len(t.rows[0].cells) >= 3:
            col2 = t.rows[0].cells[2]
            _set_cell_text(col2,
                f"Pour le Client :\n\n{v.get('client_nom','')}\n\nReprésentant légal : {v.get('client_representant','')}\n\nSignature :")

    # ── 9. Annexe A — tables conditionnelles ─────────────────────────────────
    annexe_map = {
        "site":     {"idx": 15, "fields": [
            ("site_type",          0),
            ("site_nb_pages",      1),
            (None,                 2),  # pré-rempli
            ("site_delai",         3),
            ("site_hebergement",   4),
            ("site_maintenance",   5),
            (None,                 6),  # pré-rempli
        ]},
        "seo":      {"idx": 16, "fields": [
            ("seo_zone",           0),
            ("seo_keywords",       1),
            (None,                 2),
            ("seo_rapport",        3),
            (None,                 4),
        ]},
        "ads":      {"idx": 17, "fields": [
            ("ads_plateformes",    0),
            ("ads_budget",         1),
            ("prix_ads",           2),
            ("ads_objectif",       3),
            ("ads_zone",           4),
            ("ads_rapport",        5),
        ]},
        "fidelite": {"idx": 18, "fields": [
            ("fidelite_plateforme",    0),
            (None,                     1),
            ("fidelite_maintenance",   2),
            ("fidelite_acces",         3),
        ]},
        "ia":       {"idx": 19, "fields": [
            (None,           0),
            (None,           1),
            ("ia_usecases",  2),
            ("ia_numero",    3),
            ("ia_rapport",   4),
        ]},
        "social":   {"idx": 20, "fields": [
            ("social_plateformes",  0),
            ("social_frequence",    1),
            ("social_contenus",     2),
            ("social_visuels",      3),
            ("social_validation",   4),
            ("social_rapport",      5),
        ]},
    }

    for key, cfg in annexe_map.items():
        idx = cfg["idx"]
        if key not in services:
            continue
        if idx >= len(tables):
            continue
        t = tables[idx]
        for row_i, (var_key, _) in enumerate(cfg["fields"]):
            if var_key is None:
                continue
            if row_i >= len(t.rows):
                continue
            row = t.rows[row_i]
            if len(row.cells) >= 2:
                _set_cell_text(row.cells[-1], v.get(var_key, ""))

    # ── 10. Signatures Annexe A (index 21) ───────────────────────────────────
    if len(tables) > 21:
        t = tables[21]
        if t.rows and len(t.rows[0].cells) >= 3:
            col2 = t.rows[0].cells[2]
            _set_cell_text(col2,
                f"Pour le Client :\n\n{v.get('client_nom','')}\n\nReprésentant légal : {v.get('client_representant','')}\n\nSignature :")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
